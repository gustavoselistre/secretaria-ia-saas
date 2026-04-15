"""
knowledge/parsers.py

Extração de texto plano a partir de arquivos (TXT, MD, PDF, DOCX) e URLs.
Retorna sempre uma tupla (texto, título_sugerido) para alimentar o pipeline RAG.
"""

from __future__ import annotations

import io
import logging
import os
import re
from typing import IO
from urllib.parse import urlparse

logger = logging.getLogger(__name__)

SUPPORTED_FILE_EXTENSIONS = (".txt", ".md", ".pdf", ".docx")

_WHITESPACE_RE = re.compile(r"[ \t]+")
_BLANKLINES_RE = re.compile(r"\n{3,}")


def _normalize_whitespace(text: str) -> str:
    """Colapsa espaços/tabs múltiplos e limita linhas em branco consecutivas."""
    text = _WHITESPACE_RE.sub(" ", text)
    text = _BLANKLINES_RE.sub("\n\n", text)
    return text.strip()


def _read_text_bytes(data: bytes) -> str:
    """Decodifica bytes como UTF-8, com fallback para latin-1."""
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="replace")


def _extract_txt(stream: IO[bytes]) -> str:
    return _read_text_bytes(stream.read())


def _extract_pdf(stream: IO[bytes]) -> str:
    from pypdf import PdfReader

    reader = PdfReader(stream)
    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # pypdf pode falhar em páginas corrompidas
            logger.warning("Falha ao extrair página do PDF: %s", exc)
    return "\n\n".join(p for p in pages if p.strip())


def _extract_docx(stream: IO[bytes]) -> str:
    from docx import Document

    document = Document(stream)
    paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    # Também captura texto de tabelas, que o iterador de parágrafos ignora.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def extract_text_from_file(file_obj, filename: str) -> tuple[str, str]:
    """Extrai texto plano de um arquivo enviado.

    Args:
        file_obj: File-like object com `.read()` (ex: UploadedFile do Django,
                  caminho de arquivo, BytesIO).
        filename: Nome original do arquivo — usado para detectar extensão e
                  gerar título sugerido.

    Returns:
        (texto, título_sugerido)

    Raises:
        ValueError: extensão não suportada ou arquivo sem texto extraível.
    """
    ext = os.path.splitext(filename)[1].lower()
    if ext not in SUPPORTED_FILE_EXTENSIONS:
        raise ValueError(
            f"Formato '{ext}' não suportado. "
            f"Use: {', '.join(SUPPORTED_FILE_EXTENSIONS)}."
        )

    # Aceita caminho de arquivo também
    if isinstance(file_obj, (str, bytes, os.PathLike)):
        stream: IO[bytes] = open(file_obj, "rb")
        close_after = True
    else:
        # Garante um stream de bytes seekable para pypdf/python-docx
        data = file_obj.read()
        if isinstance(data, str):
            data = data.encode("utf-8")
        stream = io.BytesIO(data)
        close_after = False

    try:
        if ext in (".txt", ".md"):
            text = _extract_txt(stream)
        elif ext == ".pdf":
            text = _extract_pdf(stream)
        elif ext == ".docx":
            text = _extract_docx(stream)
        else:  # pragma: no cover — já validado acima
            raise ValueError(f"Formato '{ext}' não suportado.")
    finally:
        if close_after:
            stream.close()

    text = _normalize_whitespace(text)
    if not text:
        raise ValueError(
            f"Nenhum texto extraível encontrado em '{filename}'. "
            "O arquivo pode estar vazio, protegido ou conter apenas imagens."
        )

    title = os.path.splitext(os.path.basename(filename))[0].strip() or filename
    return text, title


def extract_text_from_url(url: str, timeout: int = 15) -> tuple[str, str]:
    """Baixa *url* e retorna texto limpo + título.

    Remove tags de navegação/scripts/estilo. Usa BeautifulSoup.

    Raises:
        ValueError: URL inválida, resposta não-HTML, ou página sem texto.
    """
    import requests
    from bs4 import BeautifulSoup

    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https") or not parsed.netloc:
        raise ValueError(f"URL inválida: '{url}'. Use http:// ou https://.")

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (compatible; SecretariaIA/1.0; +https://secretaria-ia)"
        )
    }
    try:
        response = requests.get(url, timeout=timeout, headers=headers)
        response.raise_for_status()
    except requests.RequestException as exc:
        raise ValueError(f"Falha ao baixar '{url}': {exc}") from exc

    content_type = response.headers.get("Content-Type", "").lower()
    if "html" not in content_type and "text" not in content_type:
        raise ValueError(
            f"Content-Type '{content_type}' não suportado para URL. "
            "Baixe o arquivo manualmente e use upload de arquivo."
        )

    soup = BeautifulSoup(response.text, "html.parser")

    # Remove blocos de navegação/estilo que poluem o RAG
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
        tag.decompose()

    raw_title = soup.title.string.strip() if soup.title and soup.title.string else ""
    text = soup.get_text(separator="\n")
    text = _normalize_whitespace(text)

    if not text:
        raise ValueError(f"Nenhum texto extraível em '{url}'.")

    title = raw_title or parsed.netloc or url
    return text, title
